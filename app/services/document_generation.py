# SPDX-License-Identifier: Apache-2.0
# Copyright (c) 2025 AethOS AI

"""Generate PDF, DOCX, Markdown, and text documents under .runtime/generated_documents/."""

from __future__ import annotations

import logging
import re
import uuid
from dataclasses import dataclass, field
from datetime import date, datetime
from pathlib import Path
from typing import Any, Literal

from sqlalchemy import desc, func, select
from sqlalchemy.orm import Session

from app.models.document_artifact import DocumentArtifactModel
from app.services.handoff_paths import PROJECT_ROOT

logger = logging.getLogger(__name__)

FormatLiteral = Literal["md", "txt", "docx", "pdf"]

SUPPORTED_FORMATS: frozenset[str] = frozenset({"md", "txt", "docx", "pdf"})

_GENERIC_TITLES = frozenset(
    {
        "nexa export",
        "document",
        "export",
        "untitled",
    }
)


def get_generated_documents_base() -> Path:
    return (PROJECT_ROOT / ".runtime" / "generated_documents").resolve()


def sanitize_user_segment(user_id: str) -> str:
    s = re.sub(r"[^a-zA-Z0-9._-]+", "", (user_id or "user")[:128])
    return s or "user"


def sanitize_file_stem(title: str, *, default: str = "document") -> str:
    t = re.sub(r"[^a-zA-Z0-9._\-\s]+", " ", (title or "").strip(), flags=re.UNICODE)
    t = re.sub(r"[\s_]+", "_", t, flags=re.UNICODE).strip("._-")[:80]
    return t or default


def _unique_name(stem: str, suffix: str) -> str:
    return f"{stem}_{uuid.uuid4().hex[:10]}{suffix}"


def _slug_for_filename(s: str, max_len: int = 80) -> str:
    t = re.sub(r"[^a-zA-Z0-9._-]+", "-", (s or "").lower().strip(), flags=re.UNICODE)
    t = re.sub(r"-{2,}", "-", t).strip("-")
    t = t.replace(".", "-")
    return (t or "export")[:max_len].strip("-")


def build_smart_export_title(
    body_markdown: str,
    *,
    source_type: str = "chat",
    user_title: str | None = None,
) -> str:
    """Readable title for DB and document headings (not raw filename)."""
    u = (user_title or "").strip()
    if u and u.lower() not in _GENERIC_TITLES and len(u) >= 2:
        return u[:500]
    body = body_markdown or ""
    for line in body.splitlines():
        s = line.strip()
        if s.startswith("#"):
            inner = re.sub(r"^#+\s*", "", s).strip()
            if len(inner) >= 3:
                return inner[:500]
    plain = re.sub(r"[*_`#]+", " ", body)
    plain = re.sub(r"\s+", " ", plain).strip()
    words = plain.split()[:10] if plain else []
    if len(words) >= 3:
        return " ".join(words)[:500]
    return f"AethOS export — {source_type} — {date.today().isoformat()}"[:500]


def export_stem_for_filename(
    display_title: str,
    source_type: str = "chat",
) -> str:
    """File stem: e.g. nexa-project-plan-2026-04-26, chat-export-2026-04-26."""
    d = date.today().isoformat()
    st = (source_type or "chat").strip() or "chat"
    st_slug = _slug_for_filename(st.replace("_", "-"), max_len=32)
    core = _slug_for_filename(display_title, max_len=50)
    if st in ("chat", "memory-export", "memory_export") or st_slug in ("chat", "memory-export"):
        if core and core not in ("export", "document", "untitled", "nexa-export"):
            return f"{core}-{d}"
        return f"chat-export-{d}"
    if core and core not in ("export", "document", "untitled"):
        return f"nexa-{st_slug}-{core}-{d}"[:100]
    return f"nexa-{st_slug}-{d}"[:100]


@dataclass
class DocumentArtifact:
    id: int
    title: str
    format: str
    file_path: str
    created_at: datetime
    download_url: str | None
    metadata: dict[str, Any] = field(default_factory=dict)
    source_type: str = "chat"
    source_ref: str | None = None


class DocumentGenerationError(Exception):
    def __init__(self, message: str, code: str = "generation_failed") -> None:
        super().__init__(message)
        self.code = code


def content_looks_sensitive(text: str) -> bool:
    t = text or ""
    if re.search(r"sk-[a-zA-Z0-9]{10,}", t):
        return True
    if re.search(r"sk_ant_[a-zA-Z0-9_\-]+", t):
        return True
    if re.search(
        r"(?i)(api[_-]?key|secret|bearer|authorization)\s*[:=]\s*['\"]?[a-zA-Z0-9_\-/+=]{8,}", t
    ):
        return True
    if re.search(r"-----BEGIN (RSA |EC |OPENSSH )?PRIVATE KEY-----", t):
        return True
    return False


def _ensure_path_under_base(full: Path, base: Path) -> None:
    try:
        full_res = full.resolve()
        base_res = base.resolve()
        full_res.relative_to(base_res)
    except (OSError, ValueError) as e:
        raise DocumentGenerationError("Path escapes document root", "path_error") from e


def _write_markdown(path: Path, body: str) -> int:
    path.write_text((body or "").rstrip() + "\n", encoding="utf-8")
    return path.stat().st_size


def _write_text(path: Path, body: str) -> int:
    path.write_text((body or "").rstrip() + "\n", encoding="utf-8")
    return path.stat().st_size


def _markdownish_to_paragraphs(body: str) -> list[str]:
    t = (body or "").replace("\r\n", "\n")
    parts = re.split(r"\n{2,}", t)
    out: list[str] = []
    for p in parts:
        line = p.strip()
        if not line:
            continue
        out.append(line)
    if not out and t.strip():
        out = [t.strip()]
    return out or [""]


def _write_docx(path: Path, title: str, body: str) -> int:
    from docx import Document  # type: ignore[import-untyped]

    d = Document()
    d.add_heading((title or "Document").strip() or "Document", 0)
    for block in _markdownish_to_paragraphs(body):
        plain = re.sub(r"\*\*([^*]+)\*\*", r"\1", block)
        for line in plain.split("\n"):
            s = line.strip()
            if s.startswith("# "):
                d.add_heading(s[2:].strip(), level=1)
            elif s.startswith("## "):
                d.add_heading(s[3:].strip(), level=2)
            else:
                d.add_paragraph(line)
    d.save(str(path))
    return path.stat().st_size


def _write_pdf(path: Path, title: str, body: str) -> int:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    story: list = []
    styles = getSampleStyleSheet()
    story.append(Paragraph((title or "Document").replace("&", "&amp;"), styles["Title"]))
    story.append(Spacer(1, 12))
    for block in _markdownish_to_paragraphs(body):
        t = block.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        t = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", t)
        t = t.replace("\n", "<br/>")
        story.append(Paragraph(t, styles["BodyText"]))
        story.append(Spacer(1, 6))
    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=18,
    )
    doc.build(story)
    return path.stat().st_size


def generate_markdown_document(
    *,
    title: str,
    body_markdown: str,
    user_id: str,
    source_type: str = "chat",
    source_ref: str | None = None,
    filename_stem: str | None = None,
) -> Path:
    base = get_generated_documents_base()
    udir = base / f"user_{sanitize_user_segment(user_id)}"
    udir.mkdir(parents=True, exist_ok=True)
    stem = filename_stem or sanitize_file_stem(title)
    name = _unique_name(stem, ".md")
    out = (udir / name).resolve()
    _ensure_path_under_base(out, udir)
    _write_markdown(out, body_markdown)
    return out


def generate_text_document(
    *,
    title: str,
    body: str,
    user_id: str,
    source_type: str = "chat",
    source_ref: str | None = None,
    filename_stem: str | None = None,
) -> Path:
    base = get_generated_documents_base()
    udir = base / f"user_{sanitize_user_segment(user_id)}"
    udir.mkdir(parents=True, exist_ok=True)
    stem = filename_stem or sanitize_file_stem(title)
    name = _unique_name(stem, ".txt")
    out = (udir / name).resolve()
    _ensure_path_under_base(out, udir)
    if title.strip():
        _write_text(out, f"{title}\n\n{body}")
    else:
        _write_text(out, body)
    return out


def generate_docx_document(
    *,
    title: str,
    body_markdown: str,
    user_id: str,
    source_type: str = "chat",
    source_ref: str | None = None,
    filename_stem: str | None = None,
) -> Path:
    base = get_generated_documents_base()
    udir = base / f"user_{sanitize_user_segment(user_id)}"
    udir.mkdir(parents=True, exist_ok=True)
    stem = filename_stem or sanitize_file_stem(title)
    name = _unique_name(stem, ".docx")
    out = (udir / name).resolve()
    _ensure_path_under_base(out, udir)
    _write_docx(out, title, body_markdown)
    return out


def generate_pdf_document(
    *,
    title: str,
    body_markdown: str,
    user_id: str,
    source_type: str = "chat",
    source_ref: str | None = None,
    filename_stem: str | None = None,
) -> Path:
    base = get_generated_documents_base()
    udir = base / f"user_{sanitize_user_segment(user_id)}"
    udir.mkdir(parents=True, exist_ok=True)
    stem = filename_stem or sanitize_file_stem(title)
    name = _unique_name(stem, ".pdf")
    out = (udir / name).resolve()
    _ensure_path_under_base(out, udir)
    _write_pdf(out, title, body_markdown)
    return out


def _rel_file_path(absolute: Path) -> str:
    try:
        return str(absolute.relative_to(PROJECT_ROOT))
    except ValueError:
        return str(absolute)


def generate_document(
    db: Session,
    *,
    title: str,
    body_markdown: str,
    format: str,
    user_id: str,
    source_type: str = "chat",
    source_ref: str | None = None,
    metadata: dict[str, Any] | None = None,
    allow_sensitive: bool = False,
) -> DocumentArtifact:
    fmt = (format or "md").strip().lower()
    if fmt not in SUPPORTED_FORMATS:
        raise DocumentGenerationError(
            f"Unsupported format: {format!r} (use md, txt, docx, or pdf)", "invalid_format"
        )
    if not (body_markdown or "").strip() and not (title or "").strip():
        raise DocumentGenerationError("Nothing to export — add a title or body.", "empty_body")
    if not allow_sensitive and content_looks_sensitive(body_markdown or ""):
        raise DocumentGenerationError(
            "This document appears to contain sensitive content, so I did not export it. "
            "Remove secrets and try again.",
            "sensitive_content",
        )
    t = build_smart_export_title(
        body_markdown,
        source_type=source_type,
        user_title=(title or "").strip() or None,
    )[:500]
    file_stem = export_stem_for_filename(t, source_type=source_type)
    if fmt == "md":
        path = generate_markdown_document(
            title=t,
            body_markdown=body_markdown,
            user_id=user_id,
            source_type=source_type,
            source_ref=source_ref,
            filename_stem=file_stem,
        )
    elif fmt == "txt":
        path = generate_text_document(
            title=t,
            body=body_markdown,
            user_id=user_id,
            source_type=source_type,
            source_ref=source_ref,
            filename_stem=file_stem,
        )
    elif fmt == "docx":
        path = generate_docx_document(
            title=t,
            body_markdown=body_markdown,
            user_id=user_id,
            source_type=source_type,
            source_ref=source_ref,
            filename_stem=file_stem,
        )
    else:
        path = generate_pdf_document(
            title=t,
            body_markdown=body_markdown,
            user_id=user_id,
            source_type=source_type,
            source_ref=source_ref,
            filename_stem=file_stem,
        )

    size = path.stat().st_size
    rel = _rel_file_path(path)
    row = DocumentArtifactModel(
        owner_user_id=(user_id or "")[:64],
        title=t,
        format=fmt,
        file_path=rel,
        source_type=(source_type or "chat")[:64],
        source_ref=(source_ref or None)[:256] if source_ref else None,
    )
    meta: dict[str, Any] = {"size_bytes": size, **(metadata or {})}
    row.set_metadata(meta)
    db.add(row)
    db.commit()
    db.refresh(row)
    # Relative to /api/v1 (matches web `webFetch` paths like `/web/...`)
    durl = f"/web/documents/{row.id}/download"
    return DocumentArtifact(
        id=row.id,
        title=row.title,
        format=row.format,
        file_path=row.file_path,
        created_at=row.created_at,
        download_url=durl,
        metadata=meta,
        source_type=row.source_type,
        source_ref=row.source_ref,
    )


def list_document_artifacts_for_user(
    db: Session, owner_user_id: str, *, limit: int = 20
) -> list[DocumentArtifact]:
    rows = (
        db.execute(
            select(DocumentArtifactModel)
            .where(DocumentArtifactModel.owner_user_id == owner_user_id)
            .order_by(desc(DocumentArtifactModel.id))
            .limit(limit)
        )
        .scalars()
        .all()
    )
    out: list[DocumentArtifact] = []
    for r in rows:
        mraw = r.metadata_dict()
        meta: dict[str, Any] = mraw if isinstance(mraw, dict) else {}
        out.append(
            DocumentArtifact(
                id=r.id,
                title=r.title,
                format=r.format,
                file_path=r.file_path,
                created_at=r.created_at,
                download_url=f"/web/documents/{r.id}/download",
                metadata=meta,
                source_type=r.source_type,
                source_ref=r.source_ref,
            )
        )
    return out


def count_all_document_artifacts(db: Session) -> int:
    n = db.scalar(select(func.count()).select_from(DocumentArtifactModel))
    return int(n or 0)


def get_document_path_for_owner(db: Session, doc_id: int, owner_user_id: str) -> Path | None:
    row = db.get(DocumentArtifactModel, doc_id)
    if not row or row.owner_user_id != owner_user_id:
        return None
    p = (PROJECT_ROOT / row.file_path).resolve()
    udir = (get_generated_documents_base() / f"user_{sanitize_user_segment(owner_user_id)}").resolve()
    try:
        p.relative_to(udir)
    except ValueError:
        logger.warning("document path escape attempt id=%s path=%s", doc_id, p)
        return None
    if not p.is_file():
        return None
    return p
